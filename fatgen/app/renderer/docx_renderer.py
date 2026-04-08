"""
python-docx renderer for FAT test documents.

Builds structured Word documents following the section order:
  1. Header (company, doc number, revision, date)
  2. Title block
  3. Info table
  4. Notes (if any)
  5. Test table (test points)
  6. Step table
  7. Signature block
"""
from __future__ import annotations
import io
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ..models.project import Project
from ..models.test_config import TestCase, TestType
from ..models.tag import Tag, IOType
from ..templates.base_template import TestTemplate
from ..templates.ai_scaling import AIScalingTemplate
from ..templates.alarm_verify import AlarmVerifyTemplate
from ..templates.motor_logic import MotorLogicTemplate
from ..templates.discrete_io import DiscreteIOTemplate
from . import styles as S


def _template_for(tc: TestCase, tag: Tag, project: Project) -> TestTemplate:
    mapping = {
        TestType.AI_SCALING:   AIScalingTemplate,
        TestType.ALARM_VERIFY: AlarmVerifyTemplate,
        TestType.MOTOR_LOGIC:  MotorLogicTemplate,
        TestType.DISCRETE_IO:  DiscreteIOTemplate,
        # AO_SCALING, INTERLOCK, CUSTOM — fall back to AI for now (Phase 2)
        TestType.AO_SCALING:   AIScalingTemplate,
    }
    cls = mapping.get(tc.test_type, AIScalingTemplate)
    return cls(tc, tag, project)


class DocxRenderer:

    def render_single(self, tc: TestCase, tag: Tag, project: Project) -> bytes:
        doc = Document()
        self._setup_page(doc)
        tmpl = _template_for(tc, tag, project)
        self._build_document(doc, tmpl)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def render_all(self, test_cases: list[TestCase], project: Project) -> bytes:
        doc = Document()
        self._setup_page(doc)
        tag_map = {t.name: t for t in project.tags}

        for i, tc in enumerate(test_cases):
            tag = tag_map.get(tc.tag_name)
            if tag is None:
                continue
            if i > 0:
                self._add_page_break(doc)
            tmpl = _template_for(tc, tag, project)
            self._build_document(doc, tmpl)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ── Page setup ────────────────────────────────────────────────────────────

    def _setup_page(self, doc: Document):
        section = doc.sections[0]
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    def _add_page_break(self, doc: Document):
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_break(docx_break_type())

    # ── Document builder ──────────────────────────────────────────────────────

    def _build_document(self, doc: Document, tmpl: TestTemplate):
        meta = tmpl.get_doc_metadata()
        self._build_header(doc, meta)
        self._build_title(doc, meta)
        self._build_info_table(doc, meta)

        if meta.get("notes"):
            self._build_notes(doc, meta["notes"])
        if meta.get("sqrt_note"):
            self._build_notes(doc, meta["sqrt_note"])

        points = tmpl.get_test_points()
        if points:
            self._build_test_table(doc, points, tmpl.tc.test_type)

        steps = tmpl.get_steps()
        if steps:
            self._build_step_table(doc, steps)

        self._build_signature_block(doc, meta)

    # ── Header ────────────────────────────────────────────────────────────────

    def _build_header(self, doc: Document, meta: dict):
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        _set_table_width(tbl, Inches(7.0))

        # Left: company name
        cell_l = tbl.rows[0].cells[0]
        p = cell_l.paragraphs[0]
        run = p.add_run(meta.get("si_company", ""))
        run.font.name = "Arial"
        run.font.size = Pt(14)
        run.font.bold = True
        S.shade_cell(cell_l, "F5F5F5")

        # Centre: doc number
        cell_c = tbl.rows[0].cells[1]
        p = cell_c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(meta.get("doc_number", ""))
        run.font.name = "Arial Narrow"
        run.font.size = Pt(10)
        run.font.bold = True
        S.shade_cell(cell_c, "F5F5F5")

        # Right: revision / date
        cell_r = tbl.rows[0].cells[2]
        p = cell_r.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rev_text = f"Rev {meta.get('revision', 'A')}  |  {meta.get('revision_date', date.today().isoformat())}"
        run = p.add_run(rev_text)
        run.font.name = "Arial"
        run.font.size = Pt(9)
        S.shade_cell(cell_r, "F5F5F5")

        doc.add_paragraph()  # spacer

    # ── Title block ───────────────────────────────────────────────────────────

    def _build_title(self, doc: Document, meta: dict):
        title_text = f"FACTORY ACCEPTANCE TEST\n{meta.get('test_type','').upper()} — {meta.get('tag_name','')}"
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title_text)
        run.font.name = "Arial"
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.all_caps = True

    # ── Info table ────────────────────────────────────────────────────────────

    def _build_info_table(self, doc: Document, meta: dict):
        doc.add_paragraph()
        rows_data = [
            ("Client",          meta.get("client", "")),
            ("Site",            meta.get("client_site", "")),
            ("Project Number",  meta.get("project_number", "")),
            ("Project Name",    meta.get("project_name", "")),
            ("Tag",             meta.get("tag_name", "")),
            ("Description",     meta.get("description", "")),
            ("PLC Address",     meta.get("address", "")),
            ("Instrument",      meta.get("instrument_model", "")),
            ("Span",            meta.get("span", "")),
            ("Signal Type",     meta.get("signal_type", "")),
            ("PLC Platform",    meta.get("plc_platform", "")),
            ("PLC Model",       meta.get("plc_model", "")),
            ("Engineer",        meta.get("engineer", "")),
        ]

        tbl = doc.add_table(rows=len(rows_data), cols=4)
        tbl.style = "Table Grid"
        _set_table_width(tbl, Inches(7.0))

        for i, (key, val) in enumerate(rows_data):
            cells = tbl.rows[i].cells
            # Pair layout: key | val | key | val
            # For single-wide rows just merge right pair
            S.apply_info_key(cells[0], key)
            S.apply_table_cell(cells[1], val, center=False)
            # Leave right pair empty for now (could add second column of info)
            S.shade_cell(cells[2], "F0F0F0")
            S.shade_cell(cells[3], "FFFFFF")

        doc.add_paragraph()

    # ── Notes ─────────────────────────────────────────────────────────────────

    def _build_notes(self, doc: Document, note_text: str):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(f"NOTE: {note_text}")
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = S.NOTE_COLOR
        doc.add_paragraph()

    # ── Test point table ──────────────────────────────────────────────────────

    def _build_test_table(self, doc: Document, points: list[dict], test_type: TestType):
        self._section_heading(doc, "TEST DATA")

        if not points:
            return

        headers = list(points[0].keys())
        tbl = doc.add_table(rows=1 + len(points), cols=len(headers))
        tbl.style = "Table Grid"
        _set_table_width(tbl, Inches(7.0))

        # Header row
        for j, h in enumerate(headers):
            S.apply_table_header(tbl.rows[0].cells[j], _pretty_header(h))

        # Data rows
        for i, point in enumerate(points):
            for j, key in enumerate(headers):
                val = str(point.get(key, ""))
                cell = tbl.rows[i + 1].cells[j]
                if key in ("pass_fail", "eu_actual", "alarm_activated", "alarm_reset", "field_verified", "result"):
                    S.apply_pass_fail_cell(cell)
                else:
                    S.apply_table_cell(cell, val)

        doc.add_paragraph()

    # ── Step table ────────────────────────────────────────────────────────────

    def _build_step_table(self, doc: Document, steps: list[dict]):
        self._section_heading(doc, "TEST PROCEDURE")
        headers = ["Step", "Action", "Expected Result", "Result / Pass-Fail"]
        col_widths = [Inches(0.4), Inches(3.0), Inches(2.0), Inches(1.6)]

        tbl = doc.add_table(rows=1 + len(steps), cols=4)
        tbl.style = "Table Grid"
        _set_table_width(tbl, Inches(7.0))

        for j, h in enumerate(headers):
            S.apply_table_header(tbl.rows[0].cells[j], h, center=(j == 0))

        for i, step in enumerate(steps):
            cells = tbl.rows[i + 1].cells
            S.apply_table_cell(cells[0], str(step.get("step", "")))
            S.apply_table_cell(cells[1], step.get("action", ""), center=False)
            S.apply_table_cell(cells[2], step.get("expected", ""), center=False)
            S.apply_pass_fail_cell(cells[3])

        doc.add_paragraph()

    # ── Signature block ───────────────────────────────────────────────────────

    def _build_signature_block(self, doc: Document, meta: dict):
        self._section_heading(doc, "SIGN-OFF")
        n = min(max(meta.get("witness_blocks", 2), 1), 3)
        labels = ["SI Engineer", "Client Representative", "Instrument Technician"][:n]

        tbl = doc.add_table(rows=3, cols=n)
        tbl.style = "Table Grid"
        _set_table_width(tbl, Inches(7.0))

        row_labels = ["Name / Title", "Signature", "Date"]
        for j, label in enumerate(labels):
            # Column header
            S.apply_table_header(tbl.rows[0].cells[j], label)

        for i, row_label in enumerate(row_labels):
            for j in range(n):
                cell = tbl.rows[i].cells[j]
                if i == 0:
                    pass  # already set headers above — skip
                else:
                    S.apply_pass_fail_cell(cell)

        # Re-do first row properly (header + name row)
        # Rebuild: row0=title, row1=name, row2=sig, row3=date
        # We already have 3 rows, add a label column
        for j, label in enumerate(labels):
            S.apply_table_header(tbl.rows[0].cells[j], label)
            S.apply_info_key(tbl.rows[1].cells[j], "Name / Title")
            S.apply_info_key(tbl.rows[2].cells[j], "Signature / Date")

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _section_heading(self, doc: Document, text: str):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        _set_table_width(tbl, Inches(7.0))
        cell = tbl.rows[0].cells[0]
        cell.paragraphs[0].add_run(text)
        S.apply_section_heading(cell)
        doc.add_paragraph()  # small gap after heading — removed before table


def _set_table_width(tbl, width):
    tbl.allow_autofit = False
    tbl._tbl.tblPr.xpath('./w:tblW')[0].set(qn('w:w'), str(int(width)))
    tbl._tbl.tblPr.xpath('./w:tblW')[0].set(qn('w:type'), 'dxa')


def _pretty_header(key: str) -> str:
    replacements = {
        "pct": "% of Span",
        "mA_expected": "mA Expected",
        "vdc_expected": "VDC (250Ω)",
        "counts_expected": "ADC Counts",
        "dp_sim": "DP Simulated",
        "eu_expected": "EU Expected",
        "eu_actual": "EU Actual",
        "pass_fail": "P / F",
        "alarm_level": "Alarm",
        "setpoint": "Setpoint",
        "deadband": "Deadband",
        "plc_bit": "PLC Bit",
        "sim_value": "Sim Value",
        "reset_value": "Reset Value",
        "alarm_activated": "Activated",
        "alarm_reset": "Reset",
        "state": "State",
        "field_condition": "Field Condition",
        "plc_expected": "PLC Value",
        "hmi_expected": "HMI State",
        "field_verified": "Verified",
    }
    return replacements.get(key, key.replace("_", " ").title())


def docx_break_type():
    from docx.enum.text import WD_BREAK
    return WD_BREAK.PAGE
