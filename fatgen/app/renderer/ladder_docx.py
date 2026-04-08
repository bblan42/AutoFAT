"""
Renders LogicNetwork ladder diagrams and minimap into a python-docx Document
using tables only — no external image dependencies.

Ladder rung:  Each element occupies one cell. Wire cells connect them.
Minimap:      A single-row table with one narrow cell per rung;
              tested rungs are amber-shaded with a "▼" marker.
"""
from __future__ import annotations
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ..models.logic import LogicNetwork, LogicRung, LogicElement, ElementType
from . import styles as S

# Element symbols for table-cell rendering
_SYMBOLS = {
    "NO_CONTACT":  "─┤ ├─",
    "NC_CONTACT":  "─┤/├─",
    "COIL":        "─( )─",
    "SET_COIL":    "─(S)─",
    "RESET_COIL":  "─(R)─",
    "TON":         "┌TON┐",
    "TOF":         "┌TOF┐",
    "TP":          "┌TP ┐",
    "CTU":         "┌CTU┐",
    "CTD":         "┌CTD┐",
    "FB":          "┌ FB┐",
    "CMP":         "┌CMP┐",
}

_AMBER_FILL  = "F5A623"
_DARK_FILL   = "222222"
_MID_FILL    = "333844"
_LIGHT_FILL  = "1A1E26"
_WIRE_FILL   = "0D1218"


def render_network_to_doc(doc: Document, network: LogicNetwork):
    """Append the full network (minimap + rungs) to an existing Document."""
    _section_heading(doc, f"LOGIC — {network.name.upper()}")

    if network.description:
        p = doc.add_paragraph()
        run = p.add_run(network.description)
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x88, 0x90, 0x9A)

    # Minimap
    _render_minimap(doc, network)
    doc.add_paragraph()

    # Each rung
    for rung in network.rungs:
        _render_rung(doc, rung, network.test_rung_numbers)

    doc.add_paragraph()


def _render_minimap(doc: Document, network: LogicNetwork):
    """Single-row table — one cell per rung, amber highlight on tested rungs."""
    if not network.rungs:
        return

    n = len(network.rungs)
    tbl = doc.add_table(rows=2, cols=n)
    tbl.style = "Table Grid"

    # Row 0: rung number labels
    # Row 1: bar / indicator
    for j, rung in enumerate(network.rungs):
        is_test = rung.rung_number in network.test_rung_numbers

        # Label cell
        lbl_cell = tbl.rows[0].cells[j]
        p = lbl_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(rung.rung_number))
        run.font.name = "Arial Narrow"
        run.font.size = Pt(7)
        run.font.bold = is_test
        run.font.color.rgb = RGBColor(0xF5, 0xA6, 0x23) if is_test else RGBColor(0x60, 0x68, 0x78)
        S.shade_cell(lbl_cell, _AMBER_FILL if is_test else _MID_FILL)
        _set_cell_height(lbl_cell, Pt(12))

        # Bar cell
        bar_cell = tbl.rows[1].cells[j]
        p2 = bar_cell.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if is_test:
            run2 = p2.add_run("▼ HERE")
            run2.font.name = "Arial Narrow"
            run2.font.size = Pt(6)
            run2.font.bold = True
            run2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            S.shade_cell(bar_cell, _AMBER_FILL)
        else:
            S.shade_cell(bar_cell, _LIGHT_FILL)
        _set_cell_height(bar_cell, Pt(10))

    # Caption
    cap = doc.add_paragraph()
    cap_run = cap.add_run(
        f"Logic chain: {n} rungs total — "
        f"rung{'s' if len(network.test_rung_numbers) > 1 else ''} "
        f"{', '.join(str(r) for r in network.test_rung_numbers) or '—'} under test in this document."
    )
    cap_run.font.name = "Arial"
    cap_run.font.size = Pt(8)
    cap_run.font.italic = True
    cap_run.font.color.rgb = RGBColor(0x60, 0x68, 0x78)


def _render_rung(doc: Document, rung: LogicRung, test_rung_numbers: list[int]):
    """Render a single rung as a styled table row with element cells."""
    is_test = rung.rung_number in test_rung_numbers

    # Rung header
    hdr_tbl = doc.add_table(rows=1, cols=1)
    hdr_tbl.style = "Table Grid"
    hdr_cell = hdr_tbl.rows[0].cells[0]
    hdr_text = f"Rung {rung.rung_number}"
    if rung.description:
        hdr_text += f"  —  {rung.description}"
    if is_test:
        hdr_text += "  [▼ UNDER TEST]"
    p = hdr_cell.paragraphs[0]
    run = p.add_run(hdr_text)
    run.font.name = "Arial Narrow"
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.all_caps = True
    run.font.color.rgb = RGBColor(0xF5, 0xA6, 0x23) if is_test else RGBColor(0xEE, 0xEE, 0xEE)
    S.shade_cell(hdr_cell, "2E1E00" if is_test else _DARK_FILL)

    # Filter to main rung elements (branch_id == 0) — branches rendered separately
    main_elements = [e for e in rung.elements if e.branch_id == 0]
    branches: dict[int, list[LogicElement]] = {}
    for e in rung.elements:
        if e.branch_id > 0:
            branches.setdefault(e.branch_id, []).append(e)

    def _rung_row(elements: list[LogicElement], row_label: str = ""):
        if not elements:
            return
        # Cols: [rail] [el...] [rail]
        n_cols = len(elements) + 2
        tbl = doc.add_table(rows=2, cols=n_cols)
        tbl.style = "Table Grid"

        # Row 0: tag names above elements
        # Row 1: rail | elements | rail
        _rung_rail_cell(tbl.rows[1].cells[0])
        _rung_rail_cell(tbl.rows[1].cells[-1])
        S.shade_cell(tbl.rows[0].cells[0], _WIRE_FILL)
        S.shade_cell(tbl.rows[0].cells[-1], _WIRE_FILL)

        for j, el in enumerate(elements):
            col = j + 1
            # Top row: tag / label
            top = tbl.rows[0].cells[col]
            S.shade_cell(top, _WIRE_FILL)
            p_top = top.paragraphs[0]
            p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lbl = el.tag or el.label or el.element_type
            r = p_top.add_run(lbl)
            r.font.name = "Arial Narrow"
            r.font.size = Pt(7)
            r.font.color.rgb = RGBColor(0xF5, 0xA6, 0x23)

            # Bottom row: symbol + address
            bot = tbl.rows[1].cells[col]
            S.shade_cell(bot, _WIRE_FILL)
            p_bot = bot.paragraphs[0]
            p_bot.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sym = _SYMBOLS.get(el.element_type, f"[{el.element_type}]")
            r2 = p_bot.add_run(sym)
            r2.font.name = "Courier New"
            r2.font.size = Pt(8)
            r2.font.color.rgb = RGBColor(0x4A, 0x9E, 0xFF)
            if el.address:
                r3 = p_bot.add_run(f"\n{el.address}")
                r3.font.name = "Arial Narrow"
                r3.font.size = Pt(6)
                r3.font.color.rgb = RGBColor(0x60, 0x68, 0x78)

    _rung_row(main_elements)
    for bid, bels in branches.items():
        _rung_row(bels, f"Branch {bid}")

    doc.add_paragraph()  # spacer between rungs


def _rung_rail_cell(cell):
    S.shade_cell(cell, _DARK_FILL)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("║")
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x25, 0x2B, 0x36)


def _section_heading(doc: Document, text: str):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    cell.paragraphs[0].add_run(text)
    S.apply_section_heading(cell)
    doc.add_paragraph()


def _set_cell_height(cell, height):
    tc = cell._tc
    trPr = tc.getparent().get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height)))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)
