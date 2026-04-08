"""
python-docx style definitions for FAT documents.
All measurements in EMUs (914400 per inch) or Pt via docx helpers.
"""
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colour palette ────────────────────────────────────────────────────────────
DARK_HEADER   = RGBColor(0x22, 0x22, 0x22)   # section heading bg
MID_HEADER    = RGBColor(0x55, 0x55, 0x55)   # table header bg
LIGHT_LABEL   = RGBColor(0xEE, 0xEE, 0xEE)   # info key bg
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
BLACK         = RGBColor(0x00, 0x00, 0x00)
NOTE_COLOR    = RGBColor(0x55, 0x55, 0x55)
PASS_FAIL_BG  = RGBColor(0xF5, 0xF5, 0xF5)


def hex_to_rgb_str(color: RGBColor) -> str:
    r, g, b = color[0], color[1], color[2]
    return f"{r:02X}{g:02X}{b:02X}"


# ── Cell shading ──────────────────────────────────────────────────────────────
def shade_cell(cell, hex_color: str):
    """Apply background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, border_color: str = "AAAAAA", width_eighths: int = 4):
    """Set thin border on all sides of a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(width_eighths))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def apply_section_heading(cell):
    shade_cell(cell, hex_to_rgb_str(DARK_HEADER))
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(cell.text)
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = WHITE
    run.font.all_caps = True


def apply_table_header(cell, text: str = None, center: bool = True):
    shade_cell(cell, hex_to_rgb_str(MID_HEADER))
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.clear()
    run = p.runs[0] if p.runs else p.add_run(text or "")
    if text is not None:
        run.text = text
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = WHITE
    run.font.all_caps = True


def apply_table_cell(cell, text: str = None, center: bool = True, bold: bool = False):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text or "")
    if text is not None:
        run.text = text
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.bold = bold


def apply_info_key(cell, text: str = None):
    shade_cell(cell, hex_to_rgb_str(LIGHT_LABEL))
    apply_table_cell(cell, text=text, center=False, bold=True)


def apply_pass_fail_cell(cell):
    shade_cell(cell, hex_to_rgb_str(PASS_FAIL_BG))
    set_cell_border(cell, "AAAAAA", 4)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not p.runs:
        p.add_run("")
